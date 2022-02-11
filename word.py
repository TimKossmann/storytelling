from pydoc import doc
from random import lognormvariate
from turtle import width
from click import style
from docx import Document
from docx import shared
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
import os
import datetime
import plotly.io as pio
import plotly
import numpy as np
from docx.shared import Pt
from data_breaches_bar_chart_bubble_plot_actual_year import Charts_DataBreaches
from data_breaches_attack_vectors_treemap import Chart_AttackVectors
from phishing_graphs import Phishing_Graphs
from passwords_wordcloud import Chart_WordCloud
from docx2pdf import convert



date = datetime.date.today()
actual_year = date.strftime("%Y")

class PDF():
    def __init__(self):
        pass

    def create_pdf(self):
        #Laden der Diagramme
        #Data Breaches
        dbr = Charts_DataBreaches()
        dbr.update_bubblechart_by_year(int(actual_year) -1, False).write_image("fig_bubblechart.png")
        dbr.create_lineplot(int(actual_year) -1, False).write_image("fig_lineplot.png")
        dbr.create_table().write_image('fig_table.png', scale = 1)


        #Cyber Attacken
        atp = Chart_AttackVectors()
        atp.fig.write_image("treemap.png")
        atp.create_treemap_mensch().write_image("treemap_mensch.png")
        #Phishing
        pg = Phishing_Graphs()
        pg.get_link_donut(False, False).write_image("phishing_link.png")
        pg.get_input_donut(False, False).write_image("phishing_input.png")
        pg.get_attach_donut(False, False).write_image("phishing_attach.png")
        pg.get_fail_bar('Branche', None, False).write_image("fail_bar_mark.png",  width=900, height=800, scale=1)
        pg.get_fail_bar('Abteilung', None, False).write_image("fail_bar_type_name.png", width=900, height=800, scale=1)
        #WordCloud
        wc = Chart_WordCloud()
        wc.create_wordcloud(False).save("word_cloud.png", format="png")

        #Dokument laden
        document = Document('LaCTiS_Report - Template.docx')

        #Text-Style-Überschriften
        '''style = document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(14)'''

        # Überschriften

        header1 = 'CYBER SECURITY REPORT '
        heading = document.add_heading(header1 + '|' + actual_year, level = 0)
        header2 = document.add_paragraph().add_run('LaCTiS - Your expert in Cyber Security!')
        #header2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header2.font.size = Pt(16)
        header2.font.name = 'Calibri'
        header2.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
        header2.bold = True

        document.add_picture('LaCTis_Logo.png', width=Inches(1.5))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        header3 = document.add_paragraph().add_run('We are happy, to Cyber Secure you.')
        header3.font.size = Pt(14)
        header3.font.name = 'Calibri'
        header3.font.color.rgb = RGBColor(0x2F, 0x54, 0x96) 
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        #Einleitung
        einleitung = document.add_paragraph('Die digitale Transformation bietet den meisten Unternehmen viele neue und bedeutende Chancen. '\
        'Schnell wird jedoch klar, dass diese Transformation auch Risiken herbeiführen kann. Dabei ist fast jeder, ' \
        'also Unternehmen, öffentliche Institutionen, andere Organisationen sowie Privatpersonen das Ziel von Cyber-Attacken. ' \
        'Die Bedrohungslage nimmt stetig zu. Ziel sind die sensiblen und wertvollen Daten. '\
        'Cyber Security wird in vielen deutschen Unternehmen vernachlässigt. Oft können die Unternehmen die Risiken der Cyberkriminalität nicht abschätzen. ' \
        'Die Angst vor Cyberkriminalität ist meist nicht groß genug. '\
        'Betrachtet man im Gegenzug dazu erfasste Fälle von Cyberkriminalität in den letzten zehn Jahren in Deutschland, so kann man die '\
        'Cyberkriminalität geradezu als Wachstumsbranche bezeichnen. Neben den Hobby-Hackern gibt es mittlerweile professionelle Anbieter, die ihre Hacker-Services ' \
        'als Dienstleistung anbieten. '\
        'Besonders kleine und mittelständische Unternehmen stehen im Visier von Cyberkriminellen. Denn diese haben meist unterdurchschnittliche Sicherheitsvorkehrungen.')
        einleitung.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


        #Seitenumbruch
        document.add_page_break()

        # Thema Data Breaches
        heading2 = document.add_heading('Schaden durch Hacks', level = 1)
        heading3 = document.add_paragraph().add_run('Enstehung von Datenlecks')
        heading3.bold=True
        info_db = document.add_paragraph('Daten sind ein wertvolles Gut. Unternehmen haben oft Unmengen an Daten gespeichert, die für viele '\
            'Hacker interessant sind. Größere Plattformen, die bspw. Kreditkartendaten oder Sozialversicherungsdaten gespeichert haben, bieten '\
            'sich für Hacker als besonders attraktiv an. Datenlecks können bspw. durch Phishing-Mails, infizierte USB-Sticks, Schwachstellen von Software '\
            'oder auch Mitarbeitende enstehen. (AVG, 2019)')
        info_db.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        document.add_picture('fig_lineplot.png', width=Inches(6.0))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        abbildung_1 = document.add_paragraph('Abbildung 1: Datenlecks über die letzten 8 Jahre inklusive Durchschnitt (Quelle: information ist beautiful, 2021a)')
        abbildung_1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        info_db2 = document.add_paragraph('In den letzten Jahren ist Summe der Schäden, die durch Datenlecks in den Unternehmen entstanden sind, gestiegen (siehe dazu Abbildung 1 und Tabelle 1). '\
            'Dies hat unter anderem den Grund, dass Unternehmen immer häufiger angegriffen werden und'
            ' durch die Digitalisierung die Menge der Daten in den Unternehmen stetig steigt. Viele Daten, die vor einigen Jahren noch lokal in den Unternehmen in Papierform '\
            'vorlagen, werden heutzutage digitalisiert auf großen Servern gespeichert. Wo es Hackern früher noch gar nicht möglich war, auf Daten zuzugreifen, besteht heut für diese ein riesiges Potenzial. '\
            '')
        info_db2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY



        document.add_picture('fig_table.png', width=Inches(7.0))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table_1 = document.add_paragraph('Tabelle 1: Unternehmen mit dem höchsten Schaden durch Datenlecks von Jahr 2014 bis 2021 inklusive der Mittelwerte (Quelle: information ist beautiful, 2021a)')
        table_1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_picture('fig_bubblechart.png', width=Inches(6.0))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        abbildung_2 = document.add_paragraph('Abbildung 2: Datenlecks im Jahr 2021 (Quelle: information ist beautiful, 2021a)')
        abbildung_2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        info_db3 = document.add_paragraph('Gerade bei größeren Unternehmen, die viele Daten speichern, sind die Schäden am höchsten. Haben es die Angreifer geschafft '\
            'Zugriff zu erlangen, können sie meist direkt viele Daten abgreifen. Oft sind Unternehmen betroffen, die viele personenbezogene Daten speichern, wie bspw. Facebook und Linkedin. '\
            '(Siehe dazu Abb. 2')
        info_db3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        #Seitenumbruch
        document.add_page_break()

        #Thema Cyber Attacken
        heading2 = document.add_heading('Cyber Attacken', level = 1)
        info_ca = document.add_paragraph('Hacker haben verschiedene Angriffsziele. Zunächst denkt man dabei an Hackerangriffe auf Systeme, bei '\
            'dem der Angreifer die Schwächen des Systems ausnutzt. Allerdings können auch Menschen in das Ziel von Hackern geraten. Dabei ist das '\
            'Ziel oft, Menschen zu manipulieren oder sie bestimmte Handlungen ausführen zu lassen, sodass der Angreifer Zugang zum System bekommt.')
        info_ca.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        mensch = atp.get_percentage_human()
        info_percantage = 'Wie auch in Abbildung 3 zu erkennen liegt im Jahr ', str(int(actual_year) - 1), ' liegt die Verteilung der Cyberattacken bei ', str(mensch), ' % Mensch und ', str(100 - mensch), ' % System.' 
        info_ca1 = document.add_paragraph(info_percantage)
        info_ca1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        document.add_picture('treemap.png', width=Inches(6.0))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        abbildung_3 = document.add_paragraph('Abbildung 3: Cyber Attacken nach Angriffsvektor Mensch und System aufgeteilt (Quelle: IBM, 2021)')
        abbildung_3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        info_ca4 = document.add_paragraph('Systemschwächen sind häufig Ziele von Hackern. Dabei können alle möglichen Systeme angegriffen werden, '\
            'die mit dem Internet verbunden sind. Um mögliche Schwachstellen in Systemen zu finden, können externe Sicherheitsfirmen beauftragt werden, '\
            'die einen Angriff simuliert.')
        info_ca4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        heading = document.add_paragraph().add_run('Angriffsvektor Mensch')
        heading.bold=True
        info_ca5 = document.add_paragraph('Menschliches Fehlverhalten sind eine häufige Ursache für Datenlecks. Meist rührt das Fehlverhalten von '\
            'Unwissenheit oder Unachtsamkeit. Gegen Unachtsamkeit und mutwilliges böswilliges Verhalten kann leider schwer etwas unternommen werden. '\
            'Aber gegen Unwissenheit helfen Schulungen und Sensibilisierungskurse, die maßgeblich zum Schutz des Unternehmens beiträgt.')
        info_ca5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        
        info_ca5 = document.add_paragraph().add_run('Verteilung der Cyber Attacken mit Angriffsvektor Mensch (siehe Abb.4)')
        info_ca5.bold = True
        kom_an = atp.get_percentage_for_attack_vectors('kompromittierte Anmeldedaten')
        phish = atp.get_percentage_for_attack_vectors('Phishing')
        insider = atp.get_percentage_for_attack_vectors('böswilliger Insider')
        datenlost = atp.get_percentage_for_attack_vectors('Versehentlicher Datenverlust/verlorenes Gerät')
        mail = atp.get_percentage_for_attack_vectors('Kompromittierung von Geschäfts-E-Mails')
        document.add_paragraph('Kompromittierte Anmeldedaten: '+ str(kom_an) + ' %', style='List Bullet')
        document.add_paragraph('Phishing: '+ str(phish) + ' %', style='List Bullet')
        document.add_paragraph('Böswilliger Insider: '+ str(insider) + ' %', style='List Bullet')
        document.add_paragraph('Versehentlicher Datenverlust/verlorenes Gerät: '+ str(datenlost) + ' %', style='List Bullet')
        document.add_paragraph('Kompromittierung von Geschäfts-E-Mails: '+ str(mail) + ' %', style='List Bullet')
        document.add_picture('treemap_mensch.png', width=Inches(6.0))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        abbildung_4 = document.add_paragraph('Abbildung 4: Cyber Attacken mit Angriffsvektor Mensch (Quelle: IBM, 2021)')
        abbildung_4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info = document.add_paragraph('Im Folgenden werden die zwei größten Angriffsvektoren beim Angriffsvektor Mensch näher beschrieben.')
        info.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 
        login = document.add_paragraph().add_run('Angriffsvektor kompromittierte Anmeldedaten')
        login.bold = True
        login = atp.get_information_login()
        login_info = document.add_paragraph(login)
        login_info.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        phishing = document.add_paragraph().add_run('Angriffsvektor Phishing')
        phishing.bold = True
        phishing = atp.get_information_phishing()
        phishing_info = document.add_paragraph(phishing +' Mehr zu Phishing ist im folgenden Kapitel zu finden.' )
        phishing_info.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        


        #Seitenumbruch
        document.add_page_break()

        #Thema Phishing
        heading3 = document.add_heading('Phishing', level = 1)
        phishing_info2 = document.add_paragraph('Die meist verbreitetsten Arten von Phishing sind das Senden eines URL-Hyperlinks, die Anfrage einer Dateneingabe und ein an eine E-Mail angefügter Anhang. '\
        'Dabei sind Phishing Angriffe per URL-Hyperlink am meisten verbreitet (siehe Abb.5 - 7). Die Angreifer sind über die letzten Jahre immer erfinderischer geworden. Oft verstecken sich Angreifer hinter weitverbreiteten und vertrauenswürdigen Diensten und '\
        'täuschen so ihre Opfer, indem sie sich als echt ausgeben. Zu den meist verwendeten Themen von Phishing gehören bspw. Microsoft Teams-Anfragen, Coronavirus-Alarmmeldungen, aber auch Starbucks-Bonus oder UPS-Versandbenachrichtigungen. '\
        'Oft versuchen Angreifer ihre Opfer mit raffinierten Themen wie kostenlosen Angeboten, Rabatten oder Ticketvorverkäufen zu überzeugen. (Proofpoint, 2021)')
        phishing_info2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        document.add_picture('phishing_link.png', width=Inches(3.0))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        text_link = pg.get_text_for_dounut('link').replace(' <br>', '')
        abbildung_5 = document.add_paragraph().add_run(str(text_link))
        abbildung_5.bold=True
        abbildung_5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        abbildung_5 = document.add_paragraph('Abbildung 5: Phishing via Link (Quelle: Proofpoint, 2021)')
        abbildung_5.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_picture('phishing_input.png', width=Inches(3.0))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        text_link = pg.get_text_for_dounut('input').replace(' <br>', '')
        abbildung_6 = document.add_paragraph().add_run(str(text_link))
        abbildung_6.alignment = WD_ALIGN_PARAGRAPH.CENTER
        abbildung_6.bold=True
        abbildung_6 = document.add_paragraph('Abbildung 6: Phishing via Input (Quelle: Proofpoint, 2021)')
        abbildung_6.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_picture('phishing_attach.png', width=Inches(3.0))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        text_link = pg.get_text_for_dounut('attach').replace(' <br>', '')
        abbildung_7 = document.add_paragraph().add_run(str(text_link))
        abbildung_7.alignment = WD_ALIGN_PARAGRAPH.CENTER
        abbildung_7.bold=True
        abbildung_7 = document.add_paragraph('Abbildung 7: Phishing via Anhang (Quelle: Proofpoint, 2021)')
        abbildung_7.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mean_branche = pg.get_mean('Branche')
        mean_abteilung = pg.get_mean('Abteilung')
        phishing_info2 = document.add_paragraph('Im Folgenden werden die Fehlerquoten für 20 verschiedene Branchen und Abteilungen dargestellt. ' \
        'Die Fehlerquote bezieht sich dabei auf das Fehlverhalten bei Phishing-Angriffen durch bspw. auf Öffnen eines URL-Hyperlinks. '
        'Der Gesamtdurchschnitt über alle Branchen hinweg lag bei ' + str(mean_branche) + ' % (siehe Abb. 7). '\
        'Bei den Abteilungen lag der Durchschnitt bei ' + str(mean_abteilung) + '% (siehe Abb. 8). (Proofpoint, 2021)')
        phishing_info2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        document.add_picture('fail_bar_mark.png', width=Inches(5.0))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        abbildung_7 = document.add_paragraph('Abbildung 8: Phishing nach Branche')
        abbildung_7.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        phishing_info3 = document.add_paragraph('Fehlerquoten betrachtet auf einzelnen Abteilungsebene legt einen ausführlichen Überblick über mögliche ' \
        'Schwachstellen im Unternehmen. Angreifer haben meist einzelne Posteingänge und E-Mail-Aliase im Blick für einen Angriff. Eine Fehlerquote ' \
        'auf Unternehmensebene allein stellt nicht dar, bei welchen Teams oder Aufgabenbereiche möglicherweise Probleme auftreten.(siehe dazu Abb. 8) (Proofpoint, 2021)')
        phishing_info3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        document.add_picture('fail_bar_type_name.png', width=Inches(5.0))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        abbildung_7 = document.add_paragraph('Abbildung 9: Phishing nach Branche (Quelle: Proofpoint, 2021)')
        abbildung_7.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        #Seitenumbruch
        document.add_page_break()

        #Thema Passwortsicherheit
        heading4 = document.add_heading('Passwortsicherheit', level = 1)
        info_psw = document.add_paragraph('Ein häufiger Angriffsvektor in Bezug auf Passwörter ist das Erraten von diesen durch Ausprobieren '\
            'Dabei wird durch systematisches Ausprobieren jeglicher Kombinationen versucht, das richtige Passwort zu ermitteln. Diese Art von '\
            'Angriff wird auch Brute-Force-Angriff genannt. Hierbei werden alle Kombinationen von möglichen Passwörtern hintereinander automatisch '\
            'ausprobiert, bis das korrekte Passwort gefunden wurde. Meist führt dieses Verfahren zum Erfolg, wenn die Anzahl der möglichen Kombinationen '\
            'klein sind. So können alle Möglichkeiten in nur kurzer Zeit ausprobiert werden. Aus diesem Grund sollte die Anzahl der Zeichen so groß '\
            'wie möglich sein. Außerdem sollte das Passwort eine definierte Länge haben. So kann die Zeit bzw. Rechenleistung, die gebraucht wird, um ein Passwort '\
            'zu identifizieren, den Nutzen, das Passwort zu kennen, übersteigen. (Pohlmann, 2019)')
        info_psw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        psw_header2 =  document.add_paragraph().add_run('Verwendetes Alphabet und die Länge von Passwörtern')
        psw_header2.bold = True

        info_psw2 = document.add_paragraph('Durch das verwendete Alphabet und dabei nutzbaren Zeichen wird die Anzahl der möglichen Kombinationen bei einer bestimmten '\
            'Passwortlänge berechnet. Die Passwortlänge bezieht sich dabei auf die Anzahl der genutzten Elemente. Die Komplexität der automatischen Suche wird durch die Anzahl der möglichen '\
            'Kombinationen beschrieben. Siehe dazu Tabelle 2. (Pohlmann, 2019)')
        info_psw2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        formel_psw = document.add_paragraph('Mögliche Kombinationen = Zeichenanzahl')
        super_text = formel_psw.add_run('Passwortlänge')
        super_text.font.superscript = True
        formel_psw.alignment = WD_ALIGN_PARAGRAPH.CENTER

        psw_hinweis = document.add_paragraph('Hinweis: Es wird angenommen, dass 1 Milliarde Versuche in einer Sekunde getätigt werden können')

        data = [
            ['0-9', '10', '6', '1.000.000', '0,001 s'],
            ['', '', '8',  '100.000.000', '0,1 s'],
            ['', '', '10', '10.000.000.000', '10 s'],
            ['A-Z, a-z, 0-9', '10', '6', '56.800.235.584', '56 s'],
            ['', '', '8',  '218.340.105.584.896', '12 h'],
            ['', '', '10', '839.299.365.868.340.224', '26 Jahre'],
            ['A-Z, a-z, 0-9, ', '10', '6', '404.567.235.136', '11 min'],
            ['()[]{}?!$%&', '', '8',  '2.992.179.271.065.856', '13 h'],
            ['=*+~,.;:<>-_', '', '10', '22.130.157.888.803.070.976', '1700 Jahre']
        ]
        table_psw = document.add_table(rows = 1, cols = 5)
        table_psw.style='Medium Shading 2 Accent 5'
        header_cells = table_psw.rows[0].cells
        header_cells[0].text = 'Verwendetes\nAlphabet'
        header_cells[1].text = 'Anzahl der möglichen Zeichen'
        header_cells[2].text = 'Länge des Pass-\nworts'
        header_cells[3].text = 'Anzahl der möglichen Kombinationen'
        header_cells[4].text = 'Zeit der vollständigen Suche'

        for alphabet, anzahl, laenge, kombinationen, time in data:
            row_cells = table_psw.add_row().cells
            row_cells[0].text = alphabet.strip()
            row_cells[1].text = anzahl
            row_cells[2].text = laenge
            row_cells[3].text = kombinationen
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row_cells[4].text = time
            
            



        table_psw.cell(2,1).merge(table_psw.cell(3,1))
        table_psw.cell(1,1).merge(table_psw.cell(2,1))

        table_psw.cell(2,0).merge(table_psw.cell(3,0))  
        table_psw.cell(1,0).merge(table_psw.cell(2,0))

        table_psw.cell(5,0).merge(table_psw.cell(6,0))  
        table_psw.cell(4,0).merge(table_psw.cell(5,0))

        table_psw.cell(5,1).merge(table_psw.cell(6,1))
        table_psw.cell(4,1).merge(table_psw.cell(5,1))

        table_psw.cell(8,0).merge(table_psw.cell(9,0))  
        table_psw.cell(7,0).merge(table_psw.cell(8,0))

        table_psw.cell(8,1).merge(table_psw.cell(9,1))
        table_psw.cell(7,1).merge(table_psw.cell(8,1))

        table_2 = document.add_paragraph('Tabelle 2: Mögliche Zeichen des verwendeten alphabets und die Länge von Passwörtern (Quelle: Pohlmann, 2019)')
        table_2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        info_psw3 = document.add_paragraph('In Tabelle 2 wird deutlich, dass die Länge des Passworte und auch die Anzahl der Zeichen der verwendeten Alphabete eine '\
            'ausschlaggebende Rolle spielen')
        info_psw3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        #Seitenumbruch
        document.add_page_break()

        psw_header3 =  document.add_paragraph().add_run('Meist verwendete Passwörter')
        psw_header3.bold = True
        info_psw4 =  document.add_paragraph('Die meist verwendeten Passwörter lassen sich in elf Kategorien einteilen. Namen, machohafte Begriffe und einfache alphanumerische Zeichenketten sind '\
            'die meist verwendeten Passwortkategorien. (siehe Abb. 10) (information is beautiful, 2021b)')
        info_psw4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        document.add_picture('word_cloud.png', width=Inches(3.0))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        abbildung_8 = document.add_paragraph('Abbildung 10: Wordcloud zu den meist genutzten Passwörtern (Quelle: information is beautiful, 2021b)')
        abbildung_8.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        info_psw5 =  document.add_paragraph('Neben Datenlecks sind schlecht gewählte Passwörter die größte Sicherheitslücke. Hacker können mit Hilfe automtischer Programme tausende Zeichenkombinationen in wenigen Sekunden testen. '\
            'Ein gutes Passwort ist mindestens zehn Zeichenlang, enthält Buchstaben, Zahlen sowie Sonderzeichen und ist für jede Plattform unterschiedlich. Zur Unterstützung beim der Erstellung von Passwörtern hift die Satzregel. '\
            'Dabei wird sich ein Satz überlegt. Zum Beispiel: Mein kleiner Kater Findus spielt gerne im Garten und ist sechs Jahre alt. Für ihr Passwort nehmen sie von jedem Wort ausschließlich den ersten Buchstaben und ersetzen bspw. und mit dem &-Zeichen. '\
            'Aus dem Beispielsatz würde dann folgendes Passwort entstehen: MkKFsgiG&i6Ja. (Verbraucherzentrale, 2021)' )
        info_psw5.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY



        document.save('LaCTiS_Report.docx')
        #document.close()

        '''wdFormatPDF = 17
        inputFile = os.path.abspath("LaCTiS_Report.docx")
        outputFile = os.path.abspath("LaCTis_Report.pdf")
        word = win32com.client.Dispatch('Word.Application')
        doc = word.Documents.Open(inputFile)
        doc.SaveAs(outputFile, FileFormat=wdFormatPDF)
        doc.Close()
        word.Quit()'''
        convert('LaCTiS_Report.docx','LaCTiS_Report.pdf')

if __name__ == "__main__":
    pdf = PDF()
    pdf.create_pdf()
    
